from app.utils.PluginClass import PluginClass
from flask_jwt_extended import jwt_required, get_jwt_identity
from flask import request, send_file
from app.utils import DatabaseHandler
from app.api.records.models import RecordUpdate
from celery import shared_task, current_task
import os
import shutil
from bson.objectid import ObjectId
from app.utils.LogActions import log_actions
from app.api.logs.services import register_log
from dotenv import load_dotenv
import re
from datetime import datetime
import ffmpeg
import torch

load_dotenv()

mongodb = DatabaseHandler.DatabaseHandler()
WEB_FILES_PATH = os.environ.get('WEB_FILES_PATH', '')
USER_FILES_PATH = os.environ.get('USER_FILES_PATH', '')
ORIGINAL_FILES_PATH = os.environ.get('ORIGINAL_FILES_PATH', '')
TEMPORAL_FILES_PATH = os.environ.get('TEMPORAL_FILES_PATH', '')
HF_TOKEN = os.environ.get('HF_TOKEN', '')
batch_size = 16
compute_type='float32'

class ExtendedPluginClass(PluginClass):
    def __init__(self, path, import_name, name, description, version, author, type, settings, actions, capabilities=None, **kwargs):
        super().__init__(path, __file__, import_name, name, description, version, author, type, settings, actions = actions, capabilities = None, **kwargs)

    def add_routes(self):
        @self.route('/bulk', methods=['POST'])
        @jwt_required()
        def processing():
            current_user = get_jwt_identity()
            body = request.get_json()

            self.validate_fields(body, 'bulk')
            self.validate_roles(current_user, ['admin', 'processing'])

            task = self.bulk.delay(body, current_user)
            self.add_task_to_user(task.id, 'transcribeWhisper.bulk', current_user, 'msg')
            
            return {'msg': 'Se agregó la tarea a la fila de procesamientos'}, 201
        
        @self.route('/download', methods=['POST'])
        @jwt_required()
        def download():
            current_user = get_jwt_identity()
            body = request.get_json()
            
            self.validate_roles(current_user, ['admin', 'processing', 'editor'])

            task = self.download.delay(body, current_user)
            self.add_task_to_user(task.id, 'transcribeWhisper.download', current_user, 'file_download')
            
            return {'msg': 'Se agregó la tarea a la fila de procesamientos'}, 201
        
        @self.route('/filedownload/<taskId>', methods=['GET'])
        @jwt_required()
        def file_download(taskId):
            current_user = get_jwt_identity()

            if not self.has_role('admin', current_user) and not self.has_role('processing', current_user) and not self.has_role('editor', current_user):
                return {'msg': 'No tiene permisos suficientes'}, 401
            
            # Buscar la tarea en la base de datos
            task = mongodb.get_record('tasks', {'taskId': taskId})
            if not task:
                return {'msg': 'Tarea no existe'}, 404
            
            if task['user'] != current_user and not self.has_role('admin', current_user):
                return {'msg': 'No tiene permisos para obtener la tarea'}, 401

            if task['status'] == 'pending':
                return {'msg': 'Tarea en proceso'}, 400

            if task['status'] == 'failed':
                return {'msg': 'Tarea fallida'}, 400

            if task['status'] == 'completed':
                if task['resultType'] != 'file_download':
                    return {'msg': 'Tarea no es de tipo file_download'}, 400
                
            path = USER_FILES_PATH + task['result']
            response = send_file(path, as_attachment=True, download_name=os.path.basename(path), conditional=False)
            
            response.headers.add("Access-Control-Expose-Headers", "Content-Disposition")
            return response
        
        
    @shared_task(ignore_result=False, name='transcribeWhisper.download')
    def download(body, user):
        records_filters = {'_id': {'$in': [ObjectId(record) for record in body['records']]}}
        
        records = list(mongodb.get_all_records('records', records_filters, fields={'_id': 1, 'mime': 1, 'filepath': 1, 'processing': 1, 'name': 1, 'displayName': 1}))
        
        if len(records) == 0:
            raise Exception('No se encontraron registros')
        elif len(records) > 1:
            raise Exception('Debe seleccionar solo un registro')
        
        folder_path = USER_FILES_PATH + '/' + user + '/transcribeWhisper'
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)
        
        for r in records:
            if 'transcribeWhisper' in r['processing']:
                result = r['processing']['transcribeWhisper']['result']

                if body['format'] == 'doc':
                    from docx import Document
                    doc = Document()
                    title = r['displayName'] if 'displayName' in r else r['name']
                    doc.add_heading(title, 0)
                    doc.add_paragraph(result['text'])
                    path = os.path.join(USER_FILES_PATH, user, 'transcribeWhisper', str(r['_id']) + '.docx')
                    doc.save(path)
                    return '/' + user + '/transcribeWhisper/' + str(r['_id']) + '.docx'
                elif body['format'] == 'pdf':
                    from docx import Document
                    doc = Document()
                    title = r['displayName'] if 'displayName' in r else r['name']
                    doc.add_heading(title, 0)
                    doc.add_paragraph(result['text'])
                    temp_path = os.path.join(TEMPORAL_FILES_PATH, str(r['_id']) + '.docx')
                    doc.save(temp_path)
                    
                    try:
                        from app.plugins.filesProcessing.utils.DocumentProcessing import convert_to_pdf_with_libreoffice
                    except Exception as e:
                        raise Exception('Error al importar el módulo del plugin para el procesamiento de documentos: ' + str(e))
                    
                    output_pdf = os.path.join(USER_FILES_PATH, user, 'transcribeWhisper', str(r['_id']) + '.pdf')
                    convert_to_pdf_with_libreoffice(temp_path, output_pdf)
                    shutil.move(os.path.join(TEMPORAL_FILES_PATH, str(r['_id']) + '.pdf'), output_pdf)
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
                    return '/' + user + '/transcribeWhisper/' + str(r['_id']) + '.pdf'
                elif body['format'] == 'srt':
                    def millis_to_srt_time(millis):
                        millis = int(millis * 1000)
                        seconds, ms = divmod(millis, 1000)
                        minutes, sec = divmod(seconds, 60)
                        hours, minutes = divmod(minutes, 60)
                        return f"{hours:02d}:{minutes:02d}:{sec:02d},{ms:03d}"
                    
                    segments = result.get('segments', [])
                    srt = ''
                    for i, segment in enumerate(segments):
                        start_str = millis_to_srt_time(segment['start'])
                        end_str = millis_to_srt_time(segment['end'])
                        srt += str(i + 1) + '\n'
                        srt += start_str + ' --> ' + end_str + '\n'
                        srt += segment.get('speaker_tag', '') + segment['text'].strip() + '\n\n'
                    path = os.path.join(USER_FILES_PATH, user, 'transcribeWhisper', str(r['_id']) + '.srt')
                    with open(path, 'w', encoding='utf-8') as f:
                        f.write(srt)
                    return '/' + user + '/transcribeWhisper/' + str(r['_id']) + '.srt'
                else:
                    raise Exception('Formato no válido')
                          
    @shared_task(ignore_result=False, name='transcribeWhisper.bulk', queue='high')
    def bulk(body, user):
        current_task.update_state(state='PROGRESS', meta={
            'status': 'Iniciando procesamiento de transcripción',
            'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })
        id_process = []

        instance = ExtendedPluginClass('transcribeWhisper','', **plugin_info)
        
        if 'gpu' in body and body['gpu']:
            device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        else:
            device = torch.device('cpu')

        if 'records' not in body:
            filters = {'post_type': body['post_type']}
            if isinstance(body['post_type'], list):
                filters['post_type'] = {'$in': body['post_type']}   

            if 'parent' in body and body['parent'] and len(body['resources']) == 0:
                filters = {'$or': [{'parents.id': body['parent'], 'post_type': filters['post_type']}, {'_id': ObjectId(body['parent'])}]}
            
            if 'resources' in body and body['resources'] and len(body['resources']) > 0:
                filters = {'_id': {'$in': [ObjectId(resource) for resource in body['resources']]}, **filters}
                
            resources = list(mongodb.get_all_records('resources', filters, fields={'_id': 1}))
            resources = [str(resource['_id']) for resource in resources]

            records_filters = {
                'parent.id': {'$in': resources},
                'processing.fileProcessing': {'$exists': True},
                '$or':[{'processing.fileProcessing.type': 'audio'}, {'processing.fileProcessing.type': 'video'}]
            }
        else:
            records_filters = {'_id': {'$in': [ObjectId(record) for record in body['records']]}}

        if 'overwrite' in body and body['overwrite']:
            records_filters = {"$or": [{"processing.fileProcessing": {"$exists": False}, **records_filters}, {"processing.fileProcessing": {"$exists": True}, **records_filters}]}
        else:
            records_filters['processing.transcribeWhisper'] = {'$exists': False}
        
        records = list(mongodb.get_all_records('records', records_filters, fields={'_id': 1, 'mime': 1, 'filepath': 1, 'processing': 1}))
        
        current_task.update_state(state='PROGRESS', meta={
            'status': 'Cargando los modelos de transcripción',
            'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })

        if len(records) > 0:
            import whisper
            model = whisper.load_model(body['model'], device=device)
            
            if body['diarize']:
<<<<<<< HEAD
                from pyannote.audio import Pipeline
                diarize_model = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", token=HF_TOKEN)
=======
                import whisperx
                from whisperx import diarize
                diarize_model = diarize.DiarizationPipeline(token=HF_TOKEN, device=device)
            if body['denoise']:
                from df.enhance import enhance, init_df, load_audio, save_audio
                model_denoise, df_state, sr, _ = init_df()
>>>>>>> e83bb414c20db935968f9719ef521194d117d6f4
                
        current_task.update_state(state='PROGRESS', meta={
            'status': 'Modelos cargados, procesando registros',
            'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })

        for r in records:
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
            current_task.update_state(state='PROGRESS', meta={
                'status': 'Procesando archivo: ' + str(records.index(r) + 1) + ' de ' + str(len(records)),
                'progress': (records.index(r) + 1) / len(records) * 100,
                'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            })
            
            file_path = os.path.join(ORIGINAL_FILES_PATH, r['filepath'])
            temporal_file_path = None
            
            # 1. Native FFmpeg Extraction & Denoising
            # Always convert if not a wav, OR if denoise is explicitly requested
            if r['mime'] != 'audio/wav' or body['denoise']:
                temporal_file_path = os.path.join(TEMPORAL_FILES_PATH, str(r['_id']) + '.wav')
                os.makedirs(os.path.dirname(temporal_file_path), exist_ok=True)
                
                # Apply native FFmpeg FFT noise reduction if checked
                audio_filter = 'afftdn=nf=-25' if body['denoise'] else None
                kwargs = {'format': 'wav', 'acodec': 'pcm_s16le', 'ac': 1, 'ar': '16000', 'vn': None}
                
                if audio_filter:
                    kwargs['af'] = audio_filter
                
                try:
                    import ffmpeg
                    (
                        ffmpeg
                        .input(file_path)
                        .output(temporal_file_path, **kwargs)
                        .overwrite_output()
                        .run(capture_stdout=True, capture_stderr=True)
                    )
                except ffmpeg.Error as e:
                    error_msg = e.stderr.decode('utf8') if e.stderr else str(e)
                    print(f"--- FFMPEG CRASH LOG ---\n{error_msg}\n------------------------")
                    raise Exception('Error al extraer o limpiar el audio con FFmpeg. Revisa los logs.')
                
                file_path = temporal_file_path
            
            # 2. Transcribe with standard Whisper
            audio = whisper.load_audio(file_path)
            if body['language'] == 'auto':
                result = model.transcribe(audio)
            else:
                result = model.transcribe(audio, language=body['language'])
 
            # 3. Diarize using Pyannote (Overlap Matching logic)
            if body['diarize']:
                try:
                    current_task.update_state(state='PROGRESS', meta={
                        'status': 'Procesando segmentación del audio: ' + str(records.index(r) + 1) + ' de ' + str(len(records)),
                        'progress': (records.index(r) + 1) / len(records) * 100,
                        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    })
                    
                    # Safely load the waveform using soundfile
                    import soundfile as sf
                    data, sample_rate = sf.read(file_path)
                    waveform = torch.from_numpy(data).float()
                    if waveform.ndim == 1:
                        waveform = waveform.unsqueeze(0)
                    else:
                        waveform = waveform.transpose(0, 1)
                        
                    # Support both pyannote 3.x and 4.x wrapper outputs
                    diarization_output = diarize_model({"waveform": waveform, "sample_rate": sample_rate})
                    if hasattr(diarization_output, "speaker_diarization"):
                        diarization = diarization_output.speaker_diarization
                    else:
                        diarization = diarization_output

                    # Match Pyannote tracks to Whisper segments by finding the max overlap
                    for segment in result['segments']:
                        seg_start = segment['start']
                        seg_end = segment['end']
                        speaker_counts = {}
                        
                        for turn, _, speaker in diarization.itertracks(yield_label=True):
                            overlap_start = max(seg_start, turn.start)
                            overlap_end = min(seg_end, turn.end)
                            duration = overlap_end - overlap_start
                            
                            if duration > 0:
                                speaker_counts[speaker] = speaker_counts.get(speaker, 0) + duration
                                
                        if speaker_counts:
                            dominant_speaker = max(speaker_counts, key=speaker_counts.get)
                            segment['speaker'] = dominant_speaker.replace('SPEAKER_', 'PERSONA_')
                        else:
                            segment['speaker'] = 'PERSONA_UNKNOWN'

                except Exception as e:
                    print(f"Error en diarización: {str(e)}")
                    pass

            # 4. Reconstruct clean text and handle speaker tags
            final_text = ""
            current_speaker = None

            for segment in result['segments']:
                segment_text = segment['text']
                # Clean up hallucinated whisper text
                pattern = r'^\s*(transcribed by.*|subtitles by.*|by.*\.com|by.*\.org|http.*|.com*)$'
                if re.search(pattern, segment_text, re.IGNORECASE):
                    segment['text'] = ''
                    continue

                if body['diarize'] and 'speaker' in segment:
                    if segment['speaker'] != current_speaker:
                        current_speaker = segment['speaker']
                        # Add speaker label to final text block
                        final_text += f'\n\n{current_speaker}: {segment_text}'
                        # Tag the segment for the SRT generator to prepend
                        segment['speaker_tag'] = f"[{current_speaker}] " 
                    else:
                        final_text += ' ' + segment_text
                else:
                    final_text += ' ' + segment_text

            result['text'] = final_text.strip()

            if temporal_file_path and os.path.exists(temporal_file_path):
                os.remove(temporal_file_path)
                    
            current_task.update_state(state='PROGRESS', meta={
                'status': 'Guardando procesamiento de ' + str(records.index(r) + 1) + ' de ' + str(len(records)),
                'progress': (records.index(r) + 1) / len(records) * 100,
                'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            })

            update = {
                'processing': r['processing'],
                'updatedAt': datetime.now(),
                'updatedBy': user if user else 'system'
            }

            update['processing']['transcribeWhisper'] = {
                'type': 'av_transcribe',
                'result': result,
            }
            
            instance.update_data('records', str(r['_id']), update)
            id_process.append(r['_id'])

        register_log(user, log_actions['av_transcribe'], {'form': body, 'ids': id_process})
        instance.clear_cache()
        return f'Se procesaron {len(records)} registros'
        
general_settings = [
    {
        'type': 'checkbox',
        'label': 'Sobreescribir procesamientos existentes',
        'id': 'overwrite',
        'default': False,
        'required': False,
    },
    {
        'type': 'checkbox',
        'label': 'Limpiar audio de fondo (FFmpeg)',
        'id': 'denoise',
        'default': False,
        'required': False,
        'instructions': 'Si el audio tiene ruido de fondo estático, se filtrará usando el reductor FFT nativo de FFmpeg antes de transcribir.',
    },
    {
        'type': 'checkbox',
        'label': 'Separar parlantes (Pyannote)',
        'id': 'diarize',
        'default': False,
        'required': False,
    },
    {
        'type': 'checkbox',
        'label': 'Usar GPU (si está disponible)',
        'id': 'gpu',
        'default': False,
        'required': False,
    },
    {
        'type': 'select',
        'label': 'Tamaño del modelo Whisper',
        'id': 'model',
        'default': 'turbo',
        'options': [
            {'value': 'tiny', 'label': 'Muy pequeño'},
            {'value': 'small', 'label': 'Pequeño'},
            {'value': 'medium', 'label': 'Mediano'},
            {'value': 'large-v3', 'label': 'Grande'},
            {'value': 'turbo', 'label': 'Turbo'},
            {'value': 'large-v3-turbo', 'label': 'Turbo Grande'},
        ],
        'required': False,
    },
    {
        'type': 'select',
        'label': 'Idioma de la transcripción',
        'id': 'language',
        'default': 'auto',
        'options': [
            {'value': 'auto', 'label': 'Automático'},
            {'value': 'es', 'label': 'Español'},
            {'value': 'en', 'label': 'Inglés'},
            {'value': 'fr', 'label': 'Francés'},
            {'value': 'de', 'label': 'Alemán'},
            {'value': 'it', 'label': 'Italiano'},
            {'value': 'pt', 'label': 'Portugués'},
            
        ],
        'required': False,
    }
]
    
plugin_info = {
    'name': 'Transcripción Whisper',
    'description': 'Plugin para la transcripción automática usando Whisper y separación de audio con Pyannote.',
    'version': '0.2',
    'author': 'Néstor Andrés Peña',
    'type': ['bulk'],
    'settings': {
        'settings_bulk': [
            {
                'type':  'instructions',
                'title': 'Instrucciones',
                'text': 'Este plugin procesará todos los archivos de audio y video de los recursos hijos del recurso padre seleccionado. Utiliza Whisper para el texto y Pyannote para la separación de canales.',
            },
            *general_settings
        ]
    },
    'actions': [
        {
            'placement': 'detail_record',
            'record_type': ['audio', 'video'],
            'label': 'Transcribir con Whisper',
            'roles': ['admin', 'processing', 'editor'],
            'endpoint': 'bulk',
            'icon': 'Transcribe',
            'extraOpts': [
                *general_settings
            ]
        },
        {
            'placement': 'detail_record',
            'record_type': ['audio', 'video'],
            'label': 'Descargar transcripción',
            'roles': ['admin', 'processing', 'editor'],
            'endpoint': 'download',
            'icon': 'Download,Transcribe',
            'extraOpts': [
                {
                    'type': 'select',
                    'label': 'Formato del archivo',
                    'id': 'format',
                    'default': 'pdf',
                    'options': [
                        {'value': 'pdf', 'label': 'PDF'},
                        {'value': 'doc', 'label': 'DOC'},
                        {'value': 'srt', 'label': 'SRT'},
                    ],
                    'required': False,
                }
            ]
        }
    ]
}